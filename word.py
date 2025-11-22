import docx
from datetime import datetime, timedelta
import calendar
import os
import shutil
import sys
import re

# NOTE: This script requires the 'python-docx' library. Install it with: pip install python-docx

def copy_font_style(source_font, target_font):
    """
    Manually copies font properties from source_font to target_font
    as 'Font' objects do not have a .copy() method.
    """
    target_font.name = source_font.name
    target_font.size = source_font.size
    target_font.bold = source_font.bold
    target_font.italic = source_font.italic
    target_font.underline = source_font.underline
    target_font.strike = source_font.strike
    target_font.subscript = source_font.subscript
    target_font.superscript = source_font.superscript
    
    # Handle color, which might not be set (e.g., is None)
    try:
        if source_font.color.rgb:
            target_font.color.rgb = source_font.color.rgb
    except Exception:
        pass # Ignore color if not set or fails to copy
    
    # Handle highlight color
    try:
        if source_font.highlight_color:
             target_font.highlight_color = source_font.highlight_color
    except Exception:
        pass # Ignore highlight if not set or fails to copy

def replace_text_in_runs(paragraph, old_text, new_text):
    """
    *** ROBUST REPLACEMENT FUNCTION ***
    Finds and replaces text (old_text) within a paragraph, handling cases where 
    the text spans multiple runs (formatting blocks). It preserves formatting 
    around the replaced text.
    """
    # 1. Combine text to find the exact location of the placeholder
    full_text = "".join(run.text for run in paragraph.runs)
    
    if old_text not in full_text:
        # Also check for non-breaking space version, e.g., "Le 01/09/2025"
        old_text_nbsp = old_text.replace(" ", "\xa0")
        if old_text_nbsp not in full_text:
            return False # Text not found at all
        else:
            old_text = old_text_nbsp # Use the nbsp version

    start_index = full_text.find(old_text)
    end_index = start_index + len(old_text)

    # --- Rebuild the runs ---
    # We collect all runs, then clear the paragraph.
    # This is safer than modifying runs in-place.
    
    runs_data = []
    for run in paragraph.runs:
        # Store text, style, and font info
        runs_data.append((run.text, run.style, run.font))
    
    # Clear all runs from the paragraph
    for i in range(len(paragraph.runs)):
        r = paragraph.runs[0]
        r._r.getparent().remove(r._r)

    current_pos = 0
    replacement_done = False

    # Re-add runs, inserting the new text at the correct position
    for text, style, font in runs_data:
        run_len = len(text)
        
        # Calculate overlap
        overlap_start = max(current_pos, start_index)
        overlap_end = min(current_pos + run_len, end_index)
        
        if overlap_start < overlap_end:
            # --- This run contains part of the text to be replaced ---
            
            # 1. Add part before the old_text
            pre_text = text[:overlap_start - current_pos]
            if pre_text:
                new_run = paragraph.add_run(pre_text)
                new_run.style = style
                copy_font_style(font, new_run.font) 

            # 2. Add the new_text (only once)
            if not replacement_done:
                new_run = paragraph.add_run(new_text)
                # Apply the style of the first replaced run to the new text
                new_run.style = style
                copy_font_style(font, new_run.font) 
                replacement_done = True
            
            # 3. Add part after the old_text
            post_text = text[overlap_end - current_pos:]
            if post_text:
                new_run = paragraph.add_run(post_text)
                new_run.style = style
                copy_font_style(font, new_run.font) 
        
        else:
            # --- This run is completely outside the old_text ---
            new_run = paragraph.add_run(text)
            new_run.style = style
            copy_font_style(font, new_run.font) 
            
        current_pos += run_len
        
    return True


def process_all_text_locations(document, date_placeholder, new_date_str, name_prefix, new_name):
    """
    Helper function to run replacements in ALL text locations:
    1. Main body paragraphs
    2. Table cell paragraphs
    """
    
    # --- 1. Process main body paragraphs ---
    for paragraph in document.paragraphs:
        # A. Replace Date
        replace_text_in_runs(paragraph, date_placeholder, new_date_str)
        
        # B. Replace Name
        full_para_text = "".join(run.text for run in paragraph.runs)
        pattern = re.compile(f"({re.escape(name_prefix)}.*)") 
        match = pattern.search(full_para_text)
        if match:
            old_profile_line = match.group(1).strip()
            # If new_name is empty, the line will be "Profil : "
            new_profile_line = f"{name_prefix} {new_name}"
            replace_text_in_runs(paragraph, old_profile_line, new_profile_line)

    # --- 2. Process all paragraphs inside all tables ---
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    # A. Replace Date
                    replace_text_in_runs(paragraph, date_placeholder, new_date_str)
                    
                    # B. Replace Name
                    full_para_text = "".join(run.text for run in paragraph.runs)
                    pattern = re.compile(f"({re.escape(name_prefix)}.*)") 
                    match = pattern.search(full_para_text)
                    if match:
                        old_profile_line = match.group(1).strip()
                        # If new_name is empty, the line will be "Profil : "
                        new_profile_line = f"{name_prefix} {new_name}"
                        replace_text_in_runs(paragraph, old_profile_line, new_profile_line)


def generate_monthly_reports_docx(template_path, output_folder, target_name, target_month, target_year):
    """
    Reads a Word document template, creates copies for the specified month/year, 
    changes the date and a placeholder name in each copy, and saves the new files.
    """
    # 1. Define Placeholders and Constants
    
    # This is the exact string the script will search for in the template's text.
    DATE_PLACEHOLDER = "Le 01/09/2025" 
    
    # This searches for the prefix 'Profil :' followed by an actual name/placeholder
    # and replaces the whole line.
    NAME_PLACEHOLDER_PREFIX = "Profil\xa0:" # \xa0 is the non-breaking space
    
    # Define the starting date based on user input
    try:
        start_date = datetime(target_year, target_month, 1)
    except ValueError:
        print(f"❌ Error: Invalid month ({target_month}) or year ({target_year}).")
        return

    # Dynamically calculate the number of days in the target month
    days_in_month = calendar.monthrange(target_year, target_month)[1]

    # Format month name for display and output folder
    month_name = start_date.strftime("%B_%Y")
    
    # Update output folder based on the month and year
    final_output_folder = os.path.join(output_folder, month_name)

    # Prepare output directory
    if not os.path.exists(final_output_folder):
        os.makedirs(final_output_folder)

    print(f"--- Starting report generation for {month_name} ---")
    print(f"Targeting template: {template_path}")
    print(f"Outputting {days_in_month} files to: {final_output_folder}")
    
    try:
        # Load the template once to check structure (not used for copying)
        docx.Document(template_path)

        for i in range(days_in_month):
            # Calculate the date for the current iteration
            current_date = start_date + timedelta(days=i)
            # Format the date as required: "Le DD/MM/YYYY"
            formatted_date_fr = current_date.strftime("Le %d/%m/%Y") 
            
            # 2. Define temporary and final paths
            temp_path = os.path.join(final_output_folder, f"temp_report_{i}.docx")
            
            # Filename format: '%d-%m-%Y' (e.g., 01-11-2025)
            date_part = current_date.strftime('%d-%m-%Y') 
            
            output_filename = os.path.join(
                final_output_folder, 
                f"Rapport_d_activite_{date_part}.docx"
            )

            # 3. Create a clean copy of the template using shutil (Prevents Corruption)
            shutil.copyfile(template_path, temp_path)
            
            # 4. Open the copied document
            new_doc = docx.Document(temp_path)
            
            # --- Text Replacement Logic in ALL locations ---
            process_all_text_locations(
                new_doc,
                DATE_PLACEHOLDER,
                formatted_date_fr,
                NAME_PLACEHOLDER_PREFIX,
                TARGET_NAME
            )
                    
            # 5. Save the modified document to its final name
            new_doc.save(output_filename)
            
            # 6. Clean up the temporary file
            os.remove(temp_path)
            
            print(f"✅ Generated: {output_filename} (Date: {formatted_date_fr})")
            
    except FileNotFoundError:
        print(f"❌ Error: The template file '{template_path}' was not found. Please ensure it is in the same directory and the name is correct.")
    except Exception as e:
        print(f"❌ An unexpected error occurred: {e}")
        # Print stack trace for debugging
        import traceback
        traceback.print_exc()

# --- Configuration (EDIT THESE LINES) ---
# 1. TEMPLATE_FILE: Must be the name of your Word document.
TEMPLATE_FILE = "Template_Rapport.docx" 

# 2. OUTPUT_BASE_DIR: The base folder where the month-specific folders will be saved.
OUTPUT_BASE_DIR = "C:\\Users\\user\\Desktop\\Nouveau_dossier\\Generated_Reports" 

# 3. TARGET_NAME: The name to insert. Set to "" to keep the field empty.
TARGET_NAME = "" 
# --- End of Configuration ---

def get_user_input():
    """Prompts the user for the target month and year."""
    while True:
        try:
            # We use an arbitrary future year for the default, as 2025 was used previously.
            default_year = 2026 
            
            # Prompt for month (1-12)
            month_input = input("Enter the target month (1=Jan, 12=Dec): ")
            target_month = int(month_input)
            if not 1 <= target_month <= 12:
                raise ValueError("Month must be between 1 and 12.")

            # Prompt for year
            year_input = input(f"Enter the target year (e.g., {default_year}): ")
            target_year = int(year_input)
            if target_year < 2000 or target_year > 2100:
                raise ValueError("Year seems unrealistic. Please enter a year between 2000 and 2100.")
                
            return target_month, target_year

        except ValueError as e:
            print(f"Invalid input. Please try again: {e}")
        except Exception as e:
            print(f"An unexpected error occurred during input: {e}")

# Run the function with dynamic input
if __name__ == "__main__":
    month, year = get_user_input()
    
    # Use the base directory here; the function will create a month-specific subfolder
    generate_monthly_reports_docx(
        template_path=TEMPLATE_FILE,
        output_folder=OUTPUT_BASE_DIR,
        target_name=TARGET_NAME,
        target_month=month,
        target_year=year
    )